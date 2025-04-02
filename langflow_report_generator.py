import pandas as pd
from docx import Document
import os
from pathlib import Path
from datetime import datetime

# Component for loading Excel data
def load_excel_data_component(excel_file):
    file_path = Path(os.path.abspath("Inputs")) / excel_file
    if not file_path.exists():
        raise FileNotFoundError(f"Excel file not found at {file_path}")
    
    df = pd.read_excel(file_path)
    df.columns = [col.strip('{}') for col in df.columns]
    if 'processed' not in df.columns:
        df['processed'] = ''
    df['processed'] = df['processed'].astype(str)
    return df

# Component for loading Word template
def load_template_component(template_file):
    template_path = Path(os.path.abspath("Inputs")) / template_file
    if not template_path.exists():
        raise FileNotFoundError(f"Template file not found at {template_path}")
    
    return Document(template_path)

# Component for replacing fields in the document
def replace_fields_component(document, data_row):
    for paragraph in document.paragraphs:
        for key, value in data_row.items():
            str_value = str(value) if pd.notna(value) else ""
            placeholder = f"{{{key}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str_value)
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data_row.items():
                    str_value = str(value) if pd.notna(value) else ""
                    placeholder = f"{{{key}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str_value)
    
    return document

# Component for generating reports
def generate_reports_component(df, excel_file, template_file):
    template = load_template_component(template_file)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    processed_count = 0
    
    for index, row in df.iterrows():
        if pd.notna(row['processed']) and row['processed'] != '':
            print(f"Skipping row {index + 1} because 'processed' column is not empty ('{row['processed']}').")
            continue
        
        current_template_doc = load_template_component(template_file)
        report_doc = replace_fields_component(current_template_doc, row)
        output_filename = f"report_{timestamp}_{index + 1}.docx"
        output_path = Path(os.path.abspath("Outputs")) / output_filename
        
        report_doc.save(output_path)
        print(f"Generated report file: {output_path}")
        processed_count += 1
        
        try:
            df.loc[index, 'processed'] = output_filename
        except Exception as e:
            print(f"Error updating DataFrame in memory for row {index + 1}: {e}")

    try:
        df.to_excel(Path(os.path.abspath("Inputs")) / excel_file, index=False)
        print(f"Successfully attempted to save updates to {excel_file}")
    except PermissionError:
        print(f"\n[ERROR] Permission denied: Could not save updates to {excel_file}. Please ensure the file is closed and not open in another program, then run the script again.")
    except Exception as e:
        print(f"\n[ERROR] Failed to save updates to {excel_file}: {e}")

    return f"Processed {processed_count} new reports. Total rows in Excel: {len(df)}."

# Main component to orchestrate the workflow
def main_component():
    input_dir = Path(os.path.abspath("Inputs"))
    output_dir = Path(os.path.abspath("Outputs"))
    
    input_dir.mkdir(exist_ok=True)
    output_dir.mkdir(exist_ok=True)

    print("Select an Excel file for data:")
    excel_files = [f for f in os.listdir(input_dir) if f.endswith('.xlsx')]
    for i, file in enumerate(excel_files, 1):
        print(f"{i}. {file}")
    excel_file = excel_files[int(input("Select a file (1-{}): ".format(len(excel_files)))) - 1]

    print("\nSelect a Word template file:")
    template_files = [f for f in os.listdir(input_dir) if f.endswith('.docx')]
    for i, file in enumerate(template_files, 1):
        print(f"{i}. {file}")
    template_file = template_files[int(input("Select a file (1-{}): ".format(len(template_files)))) - 1]

    print(f"\nSelected Excel file: {excel_file}")
    print(f"Selected template: {template_file}")
    confirm = input("Proceed with these files? (y/n): ").lower()
    
    if confirm == 'y':
        df = load_excel_data_component(excel_file)
        result = generate_reports_component(df, excel_file, template_file)
        print(result)
    else:
        print("Operation cancelled.")

if __name__ == "__main__":
    main_component()
