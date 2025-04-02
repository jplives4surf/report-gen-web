import pandas as pd
from docx import Document
import os
from pathlib import Path
from datetime import datetime

class ReportGenerator:
    def __init__(self, input_dir="Inputs", output_dir="Outputs"):
        self.input_dir = Path(os.path.abspath(input_dir))
        self.output_dir = Path(os.path.abspath(output_dir))
        
        self.input_dir.mkdir(exist_ok=True)
        self.output_dir.mkdir(exist_ok=True)

    def load_excel_data(self, excel_file):
        file_path = self.input_dir / excel_file
        if not file_path.exists():
            raise FileNotFoundError(f"Excel file not found at {file_path}")
        
        df = pd.read_excel(file_path)
        df.columns = [col.strip('{}') for col in df.columns]
        if 'processed' not in df.columns:
            df['processed'] = ''
        # Ensure the 'processed' column is treated as string type to avoid dtype warnings
        df['processed'] = df['processed'].astype(str) 
        return df

    def load_template(self, template_file):
        template_path = self.input_dir / template_file
        if not template_path.exists():
            raise FileNotFoundError(f"Template file not found at {template_path}")
        
        return Document(template_path)

    def replace_fields(self, document, data_row):
        for paragraph in document.paragraphs:
            for key, value in data_row.items():
                str_value = str(value) if pd.notna(value) else ""
                placeholder = f"{{{key}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str_value)
        
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data_row.items():
                        str_value = str(value) if pd.notna(value) else ""
                        placeholder = f"{{{key}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str_value)
        
        return document

    def generate_reports(self, df, excel_file, template_file):
        template = self.load_template(template_file)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        processed_count = 0 # Keep track of newly processed reports
        
        for index, row in df.iterrows():
            # More robust check for already processed rows (handles NaN and empty strings)
            if pd.notna(row['processed']) and row['processed'] != '': 
                print(f"Skipping row {index + 1} because 'processed' column is not empty ('{row['processed']}').")
                continue
            
            # Create a fresh template instance for each report to avoid cumulative changes
            current_template_doc = self.load_template(template_file) 
            report_doc = self.replace_fields(current_template_doc, row)
            output_filename = f"report_{timestamp}_{index + 1}.docx"
            output_path = self.output_dir / output_filename
            
            report_doc.save(output_path)
            print(f"Generated report file: {output_path}") # Console output includes filename
            processed_count += 1
            
            try:
                # Update 'processed' column with the generated filename
                df.loc[index, 'processed'] = output_filename 
                # print(f"DEBUG: Set df.loc[{index}, 'processed'] = {output_filename}") # Optional debug print
            except Exception as e: # Catch potential errors during DataFrame update
                 print(f"Error updating DataFrame in memory for row {index + 1}: {e}")

        
        try:
            # Attempt to save all changes back to the Excel file
            df.to_excel(self.input_dir / excel_file, index=False)
            print(f"Successfully attempted to save updates to {excel_file}") 
        except PermissionError:
            # Explicit message if saving fails due to permissions
            print(f"\n[ERROR] Permission denied: Could not save updates to {excel_file}. Please ensure the file is closed and not open in another program, then run the script again.")
        except Exception as e:
            # Catch other potential saving errors
            print(f"\n[ERROR] Failed to save updates to {excel_file}: {e}")

            
        # Return message reflects newly generated reports
        return f"Processed {processed_count} new reports. Total rows in Excel: {len(df)}."

def get_file_selection(directory, extension):
    files = [f for f in os.listdir(directory) if f.endswith(extension)]
    if not files:
        return None
    
    print(f"\nAvailable {extension} files:")
    for i, file in enumerate(files, 1):
        print(f"{i}. {file}")
    
    while True:
        try:
            choice = int(input(f"Select a file (1-{len(files)}): "))
            if 1 <= choice <= len(files):
                return files[choice - 1]
            print("Invalid selection. Try again.")
        except ValueError:
            print("Please enter a number.")

def main():
    try:
        generator = ReportGenerator()
        
        print("Select an Excel file for data:")
        excel_file = get_file_selection(generator.input_dir, '.xlsx')
        if not excel_file:
            raise FileNotFoundError("No Excel files found in Inputs directory")
        
        print("\nSelect a Word template file:")
        template_file = get_file_selection(generator.input_dir, '.docx')
        if not template_file:
            raise FileNotFoundError("No Word template files found in Inputs directory")
        
        print(f"\nSelected Excel file: {excel_file}")
        print(f"Selected template: {template_file}")
        confirm = input("Proceed with these files? (y/n): ").lower()
        
        if confirm == 'y':
            df = generator.load_excel_data(excel_file)
            result = generator.generate_reports(df, excel_file, template_file)
            print(result)
        else:
            print("Operation cancelled.")
            
    except Exception as e:
        print(f"Error: {str(e)}")

if __name__ == "__main__":
    main()
