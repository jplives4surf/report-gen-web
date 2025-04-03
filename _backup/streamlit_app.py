import streamlit as st
import pandas as pd
from docx import Document
import io
import zipfile
from datetime import datetime
import os
import re

# Function to replace fields in the document (adapted from report_generator.py)
def replace_fields(document, data_row):
    """Replaces placeholders in paragraphs and tables of a docx document."""
    # Process paragraphs
    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        
        # Try to find placeholders in the paragraph text
        # This regex will match {{variable}} with or without spaces inside the braces
        placeholders = re.findall(r'\{\{\s*([^}]+)\s*\}\}', paragraph_text)
        
        if placeholders:
            # Make a copy of the original text
            new_text = paragraph_text
            
            for placeholder in placeholders:
                # Clean up the placeholder name
                clean_placeholder = placeholder.strip()
                
                # Check if the placeholder exists in the data_row
                for key, value in data_row.items():
                    str_key = str(key).strip()
                    # Check if the key matches the placeholder (case-insensitive)
                    if str_key.lower() == clean_placeholder.lower():
                        if pd.notna(value):
                            if isinstance(value, datetime):
                                str_value = value.strftime('%Y-%m-%d')  # Format date without time
                            else:
                                str_value = str(value)
                        else:
                            str_value = ""
                        
                        # Replace in the paragraph text
                        full_placeholder = f"{{{{{placeholder}}}}}"
                        new_text = new_text.replace(full_placeholder, str_value)
            
            # Set the paragraph text to the new text
            if new_text != paragraph_text:
                # Clear the paragraph
                p = paragraph._p
                p.clear_content()
                
                # Add a new run with the new text
                paragraph.add_run(new_text)

    # Process tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_text = paragraph.text
                    
                    # Try to find placeholders in the paragraph text
                    placeholders = re.findall(r'\{\{\s*([^}]+)\s*\}\}', paragraph_text)
                    
                    if placeholders:
                        # Make a copy of the original text
                        new_text = paragraph_text
                        
                        for placeholder in placeholders:
                            # Clean up the placeholder name
                            clean_placeholder = placeholder.strip()
                            
                            # Check if the placeholder exists in the data_row
                            for key, value in data_row.items():
                                str_key = str(key).strip()
                                
                                # Check if the key matches the placeholder (case-insensitive)
                                if str_key.lower() == clean_placeholder.lower():
                                    str_value = str(value) if pd.notna(value) else ""
                                    
                                    # Replace in the paragraph text
                                    full_placeholder = f"{{{{{placeholder}}}}}"
                                    new_text = new_text.replace(full_placeholder, str_value)
                        
                        # Set the paragraph text to the new text
                        if new_text != paragraph_text:
                            # Clear the paragraph
                            p = paragraph._p
                            p.clear_content()
                            
                            # Add a new run with the new text
                            paragraph.add_run(new_text)
    
    return document

# Function to reset the app state
def reset_app():
    """Reset the app state by clearing all session state variables."""
    # Clear generated reports
    st.session_state.generated_reports = {}
    st.session_state.generated_zip_bytes = None
    st.session_state.generated_zip_filename = None
    st.session_state.processed_count = 0
    st.session_state.skipped_count = 0
    st.session_state.total_rows = 0
    st.session_state.excel_data = None
    st.session_state.updated_excel_bytes = None
    st.session_state.updated_excel_filename = None
    
    # Rerun the app to refresh the UI
    st.rerun()  # Use st.rerun() instead of st.experimental_rerun()

# --- Streamlit App ---
st.set_page_config(layout="wide")
st.title("📄 Report Generator & Intake Creator")

# --- Tabs ---
tab1, tab2 = st.tabs(["Create Report", "Create Intake XLS"])

# --- Tab 1: Create Report ---
with tab1:
    st.header("Create Report from Template and Data")
    st.info("Upload an Excel file and a Word template. Reports will be generated for all rows where the 'processed' column (if exists) is empty. The generated reports will be available for individual download, and can also be downloaded as a zip file. The Excel file will be updated with the filename of the generated report in the 'processed' column.")

    # Add note about placeholder format
    st.warning("Note: In your Word template, use double curly braces for placeholders: {{variable_name}}")

    # Add a reset button specific to this tab's state
    if st.button("Reset Report Generation", key="reset_report"):
        reset_app() # Assuming reset_app clears relevant state for report gen

    # File Uploaders
    col1, col2 = st.columns(2)
    with col1:
        uploaded_excel = st.file_uploader("1. Upload Excel Intake File (.xlsx)", type="xlsx", key="uploaded_excel")
    with col2:
        uploaded_template = st.file_uploader("2. Upload Word Template File (.docx)", type="docx", key="uploaded_template")

    # State for generated reports and zip file
    if 'generated_reports' not in st.session_state:
        st.session_state.generated_reports = {}  # {filename: bytes}
    if 'generated_zip_bytes' not in st.session_state:
        st.session_state.generated_zip_bytes = None
    if 'generated_zip_filename' not in st.session_state:
        st.session_state.generated_zip_filename = None
    if 'processed_count' not in st.session_state:
        st.session_state.processed_count = 0
    if 'skipped_count' not in st.session_state:
        st.session_state.skipped_count = 0
    if 'total_rows' not in st.session_state:
        st.session_state.total_rows = 0
    if 'excel_data' not in st.session_state:
        st.session_state.excel_data = None
    if 'updated_excel_bytes' not in st.session_state:
        st.session_state.updated_excel_bytes = None
    if 'updated_excel_filename' not in st.session_state:
        st.session_state.updated_excel_filename = None


    if uploaded_excel is not None and uploaded_template is not None:
        if st.button("Generate Reports"):
            # Reset state
            st.session_state.generated_reports = {}
            st.session_state.generated_zip_bytes = None
            st.session_state.generated_zip_filename = None
            st.session_state.processed_count = 0
            st.session_state.skipped_count = 0
            st.session_state.total_rows = 0
            st.session_state.excel_data = None
            st.session_state.updated_excel_bytes = None
            st.session_state.updated_excel_filename = None

            # Create a status area that will be updated during processing
            status_area = st.empty()
            status_area.info("Starting report generation process...")
            
            # Create a log area for detailed messages
            log_container = st.container()
            with log_container:
                st.subheader("Processing Log")
                log_area = st.empty()
                log_messages = []
                
                def add_log(message):
                    log_messages.append(f"{datetime.now().strftime('%H:%M:%S')} - {message}")
                    log_area.code("\n".join(log_messages), language="")

            try:
                # Load Excel data
                add_log(f"Loading Excel file: {uploaded_excel.name}")
                status_area.info("Loading Excel data...")
                excel_bytes = io.BytesIO(uploaded_excel.getvalue())
                df = pd.read_excel(excel_bytes)
                st.session_state.total_rows = len(df)
                add_log(f"Found {len(df)} rows in Excel file")

                # Get the original filename
                original_filename = uploaded_excel.name

                # Clean column names
                add_log("Cleaning column names")
                df.columns = [str(col).strip('{}') for col in df.columns]
                add_log(f"Columns: {', '.join(df.columns.tolist())}")

                # Check for 'processed' column, add if missing
                if 'processed' not in df.columns:
                    add_log("Adding 'processed' column (not found in Excel)")
                    df['processed'] = ''
                else:
                    add_log("Found existing 'processed' column")
                    
                # Ensure 'processed' is string type to handle various empty values consistently
                df['processed'] = df['processed'].fillna('').astype(str)

                # Store the original excel data in session state
                st.session_state.excel_data = df.copy()

                # Load template bytes once
                add_log(f"Loading Word template: {uploaded_template.name}")
                status_area.info("Loading Word template...")
                template_bytes = io.BytesIO(uploaded_template.getvalue())

                # Generate individual reports and store them
                timestamp_run = datetime.now().strftime("%Y%m%d_%H%M%S")
                add_log(f"Starting report generation with timestamp: {timestamp_run}")
                
                # Create a progress bar
                progress_bar = st.progress(0)
                
                for index, row in df.iterrows():
                    # Update status with current row
                    progress_percent = int((index / len(df)) * 100)
                    progress_bar.progress(progress_percent)
                    status_area.info(f"Processing row {index + 1} of {len(df)} ({progress_percent}%)...")
                    
                    # Check if row should be processed
                    if row['processed'] != '':
                        add_log(f"Row {index + 1}: Skipped (already processed as '{row['processed']}')")
                        st.session_state.skipped_count += 1
                        continue  # Skip if 'processed' column is not empty

                    # Show what data we're processing
                    row_data = row.to_dict()
                    sample_data = {k: v for i, (k, v) in enumerate(row_data.items()) if i < 3}  # Show first 3 fields
                    add_log(f"Row {index + 1}: Processing data {sample_data}...")
                    
                    # Load a fresh template instance for each report
                    template_bytes.seek(0)  # Reset stream position
                    document = Document(template_bytes)
                    add_log(f"Row {index + 1}: Replacing fields in template")
                    report_doc = replace_fields(document, row.to_dict())

                    # Save the generated document to a temporary byte stream
                    doc_buffer = io.BytesIO()
                    report_doc.save(doc_buffer)
                    doc_buffer.seek(0)

                    # Store the generated report
                    output_filename = f"report_{timestamp_run}_{index + 1}.docx"
                    add_log(f"Row {index + 1}: Generated report '{output_filename}'")
                    st.session_state.generated_reports[output_filename] = doc_buffer.getvalue()

                    # Update the 'processed' column in the DataFrame
                    st.session_state.excel_data.loc[index, 'processed'] = output_filename
                    st.session_state.processed_count += 1

                # Complete the progress bar
                progress_bar.progress(100)
                
                if st.session_state.processed_count > 0:
                    status_message = f"Generated {st.session_state.processed_count} reports. Skipped {st.session_state.skipped_count} previously processed rows (out of {st.session_state.total_rows} total)."
                    status_area.success(status_message)
                    add_log(status_message)
                else:
                    status_message = f"No new reports generated. All {st.session_state.total_rows} rows were already marked as processed or the file was empty."
                    status_area.warning(status_message)
                    add_log(status_message)

                # Create zip file (after individual reports are generated)
                if st.session_state.generated_reports:
                    add_log("Creating ZIP file with all generated reports")
                    status_area.info("Creating ZIP file...")
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                        for filename, report_bytes in st.session_state.generated_reports.items():
                            zip_file.writestr(filename, report_bytes)
                    st.session_state.generated_zip_bytes = zip_buffer.getvalue()
                    st.session_state.generated_zip_filename = f"generated_reports_{timestamp_run}.zip"
                    add_log(f"ZIP file created: {st.session_state.generated_zip_filename}")

                # Save the updated DataFrame to a new Excel file
                add_log("Updating Excel file with processing status")
                status_area.info("Saving updated Excel file...")
                excel_bytes_output = io.BytesIO()
                st.session_state.excel_data.to_excel(excel_bytes_output, index=False)
                excel_bytes_output.seek(0)
                st.session_state.updated_excel_bytes = excel_bytes_output.getvalue()
                
                # Create a filename for the updated Excel file
                filename_parts = original_filename.rsplit('.', 1)
                base_name = filename_parts[0]
                extension = filename_parts[1] if len(filename_parts) > 1 else 'xlsx'
                st.session_state.updated_excel_filename = f"{base_name}_updated_{timestamp_run}.{extension}"
                add_log(f"Updated Excel file created: {st.session_state.updated_excel_filename}")
                
                # Final status update
                status_area.success("Processing complete! You can download the generated files below.")

            except Exception as e:
                st.error(f"An error occurred during report generation: {e}")
                # Reset state
                st.session_state.generated_reports = {}
                st.session_state.generated_zip_bytes = None
                st.session_state.generated_zip_filename = None
                st.session_state.processed_count = 0
                st.session_state.skipped_count = 0
                st.session_state.total_rows = 0
                st.session_state.excel_data = None
                st.session_state.updated_excel_bytes = None
                st.session_state.updated_excel_filename = None

    # Display individual download buttons
    if st.session_state.generated_reports:
        st.subheader("Download Individual Reports:")
        for filename, report_bytes in st.session_state.generated_reports.items():
            st.download_button(
                label=f"⬇️ Download {filename}",
                data=report_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_{filename}"  # Unique key for each button
            )

    # Display Zip Download Button
    if st.session_state.generated_zip_bytes:
        st.subheader("Download All Reports as ZIP:")
        st.download_button(
            label=f"⬇️ Download All Reports ({len(st.session_state.generated_reports)} files) as ZIP",
            data=st.session_state.generated_zip_bytes,
            file_name=st.session_state.generated_zip_filename,
            mime="application/zip",
            key="download_zip"
        )

    # Display Updated Excel Download Button
    if st.session_state.updated_excel_bytes:
        st.subheader("Download Updated Excel File:")
        st.download_button(
            label=f"⬇️ Download Updated Excel File (with 'processed' column updated)",
            data=st.session_state.updated_excel_bytes,
            file_name=st.session_state.updated_excel_filename,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="download_excel"
        )

    # Display message if all rows were skipped
    if st.session_state.total_rows > 0 and st.session_state.processed_count == 0 and st.session_state.skipped_count == st.session_state.total_rows:
        # Explicit message if all rows were skipped
        st.info(f"All {st.session_state.total_rows} rows were already marked as processed. No new reports generated.")


# --- Tab 2: Create Intake XLS ---
with tab2:
    st.header("Create Intake Excel Template")
    st.info("This section allows you to create a blank Excel template based on the placeholders found in a Word document.")

    uploaded_template_for_xls = st.file_uploader("Upload Word Template to Extract Placeholders (.docx)", type="docx", key="uploaded_template_xls")

    if uploaded_template_for_xls is not None:
        if st.button("Generate Intake Template", key="generate_xls"):
            try:
                st.info("Processing template to find placeholders...")
                template_bytes_xls = io.BytesIO(uploaded_template_for_xls.getvalue())
                document_xls = Document(template_bytes_xls)
                placeholders = set() # Use a set to avoid duplicates

                # Function to extract placeholders (can be reused or adapted)
                def extract_placeholders(text):
                    return re.findall(r'\[\s*([^\]]+)\s*\]', text)

                # Extract from paragraphs
                for paragraph in document_xls.paragraphs:
                    found = extract_placeholders(paragraph.text)
                    for p in found:
                        placeholders.add(p.strip())

                # Extract from tables
                for table in document_xls.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                found = extract_placeholders(paragraph.text)
                                for p in found:
                                    placeholders.add(p.strip())

                if not placeholders:
                    st.warning("No placeholders found in the format {{variable_name}}.")
                else:
                    st.success(f"Found {len(placeholders)} unique placeholders: {', '.join(sorted(list(placeholders)))}")

                    # Create an Excel file
                    output_df = pd.DataFrame(columns=sorted(list(placeholders)))
                    # Add the 'processed' column, essential for the report generation tab
                    if 'processed' not in output_df.columns:
                         output_df['processed'] = None # Add it if not extracted

                    excel_output_buffer = io.BytesIO()
                    output_df.to_excel(excel_output_buffer, index=False)
                    excel_output_buffer.seek(0)

                    xls_filename = f"intake_template_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

                    st.download_button(
                        label=f"⬇️ Download Intake Template ({xls_filename})",
                        data=excel_output_buffer,
                        file_name=xls_filename,
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="download_intake_xls"
                    )

            except Exception as e:
                st.error(f"An error occurred while generating the intake template: {e}")
